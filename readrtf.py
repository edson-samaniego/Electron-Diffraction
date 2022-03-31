import cv2
import docx #(pip install --pre python-docx)
import aspose.words as aw#(pip install aspose-words)
                         #convert rtf to word

doc = aw.Document("Au.rtf")
doc.save("Au.docx")


def readoc(file):
    doc = docx.Document(file)
    completedText=[]

    for paragraph in doc.paragraphs:
        completedText.append(paragraph.text)
    return '\n' .join(completedText)

tarjeta=readoc('Au.docx')
print(tarjeta)
doc = docx.Document('Au.docx')
renglones=[]
for s in doc.paragraphs:
    renglones.append(s.text)
if 'Peak list' in renglones:
    print('si lo encontro')
    pos=renglones.index('Peak list')
    posfinal=renglones.index('Structure')
    print('posicion',pos)
    tabla=renglones[pos:posfinal]

for i in tabla:
    print(i)


    
